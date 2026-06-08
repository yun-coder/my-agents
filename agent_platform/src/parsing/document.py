"""多格式文档解析器：支持 PDF、Word、Markdown、TXT、HTML。

使用 PyMuPDF (fitz) 解析 PDF，python-docx 解析 Word，docling 作为备用。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".doc", ".md", ".txt", ".html", ".htm"}


@dataclass
class ParsedDocument:
    """解析后的文档，保留元数据。"""

    file_path: str
    text: str
    metadata: dict = field(default_factory=dict)
    page_count: int = 0


@dataclass
class TextChunk:
    """文本分块，带元数据用于引用来源。"""

    text: str
    source: str  # 文件路径 + 分块ID
    chunk_index: int
    metadata: dict = field(default_factory=dict)


def parse_pdf_pymupdf(file_path: Path) -> ParsedDocument:
    """使用 PyMuPDF 解析 PDF（开源，无需 Key）。"""
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    pages: list[str] = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()

    return ParsedDocument(
        file_path=str(file_path),
        text="\n\n".join(pages),
        metadata={"format": "pdf", "parser": "pymupdf"},
        page_count=len(pages),
    )


def parse_docx(file_path: Path) -> ParsedDocument:
    """使用 python-docx 解析 Word 文档。"""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return ParsedDocument(
        file_path=str(file_path),
        text="\n\n".join(paragraphs),
        metadata={"format": "docx", "parser": "python-docx"},
    )


def parse_text_file(file_path: Path) -> ParsedDocument:
    """解析纯文本文件（md, txt, html）。"""
    text = file_path.read_text(encoding="utf-8")
    suffix = file_path.suffix.lower()
    return ParsedDocument(
        file_path=str(file_path),
        text=text,
        metadata={"format": suffix.lstrip("."), "parser": "builtin"},
    )


PARSERS: dict[str, Callable[[Path], ParsedDocument]] = {
    ".pdf": parse_pdf_pymupdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".md": parse_text_file,
    ".txt": parse_text_file,
    ".html": parse_text_file,
    ".htm": parse_text_file,
}


def parse_document(file_path: str | Path) -> ParsedDocument:
    """自动识别文件类型并解析。"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    suffix = path.suffix.lower()
    if suffix not in PARSERS:
        raise ValueError(
            f"不支持的文件格式：{suffix}。支持：{sorted(SUPPORTED_SUFFIXES)}"
        )
    logger.info("解析文档：%s（格式：%s）", path.name, suffix)
    return PARSERS[suffix](path)


def chunk_text(
    text: str,
    *,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    source: str = "unknown",
) -> list[TextChunk]:
    """简单的滑动窗口分块器。

    生产环境可替换为 LangChain RecursiveCharacterTextSplitter。
    这里提供纯 Python 实现，零依赖。
    """
    if not text.strip():
        return []

    paragraphs = text.split("\n")
    chunks: list[TextChunk] = []
    current_chunk = ""
    chunk_index = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current_chunk) + len(para) + 1 <= chunk_size:
            current_chunk = (current_chunk + "\n" + para).strip()
        else:
            if current_chunk:
                chunks.append(
                    TextChunk(
                        text=current_chunk,
                        source=f"{source}#chunk{chunk_index}",
                        chunk_index=chunk_index,
                    )
                )
                chunk_index += 1
                overlap_text = (
                    current_chunk[-chunk_overlap:] if chunk_overlap > 0 else ""
                )
                current_chunk = (overlap_text + "\n" + para).strip() if overlap_text else para
    if current_chunk:
        chunks.append(
            TextChunk(
                text=current_chunk,
                source=f"{source}#chunk{chunk_index}",
                chunk_index=chunk_index,
            )
        )
    return chunks


def parse_and_chunk(
    file_path: str | Path,
    *,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> list[TextChunk]:
    """解析文档并分块，一步完成。"""
    doc = parse_document(file_path)
    return chunk_text(
        doc.text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        source=Path(file_path).name,
    )
